import os
import io
import pandas as pd
from docx import Document
from docx2pdf import convert
from PyPDF2 import PdfReader, PdfWriter

print("=" * 60)
print(" MIC LETTER GENERATION SCRIPT ")
print("=" * 60)

# =====================================================
# CONFIGURATION & PATHS
# =====================================================
CSV_FILE = "members.csv"
DEFAULT_TEMPLATE = "MIC_Letter_Normal.docx"
DOCX_OUTPUT_DIR = "generated_docx"
PDF_OUTPUT_DIR = "generated_pdf"

os.makedirs(DOCX_OUTPUT_DIR, exist_ok=True)
os.makedirs(PDF_OUTPUT_DIR, exist_ok=True)

# =====================================================
# READ CSV DATA
# =====================================================
if not os.path.exists(CSV_FILE):
    # Fallback check
    if os.path.exists("members.xlsx"):
        df = pd.read_excel("members.xlsx")
    else:
        raise FileNotFoundError(f"Data file '{CSV_FILE}' not found!")
else:
    try:
        df = pd.read_csv(CSV_FILE, encoding="utf-8-sig")
    except UnicodeDecodeError:
        df = pd.read_csv(CSV_FILE, encoding="latin1")

# Clean column headers (strip spaces, remove curly brackets {{ }} if present)
raw_columns = df.columns.tolist()
cleaned_col_map = {}
for col in raw_columns:
    clean_name = str(col).strip().replace("{{", "").replace("}}", "").strip()
    cleaned_col_map[col] = clean_name

df.rename(columns=cleaned_col_map, inplace=True)

print(f"\nLoaded columns: {list(df.columns)}")
print(f"Total rows to process: {len(df)}")

# Helper to find column case-insensitively
def get_column_value(row, possible_names, default=""):
    lower_map = {str(k).lower(): k for k in row.index}
    for name in possible_names:
        name_lower = name.lower()
        if name_lower in lower_map:
            val = row[lower_map[name_lower]]
            if pd.notna(val):
                return str(val).strip()
    return default

# =====================================================
# HEAD & POSITION MAPPINGS
# =====================================================
dept_heads = {
    "Management": {
        "head1": "Ramakrishnan P H",
        "head1_post": "Management Head",
        "head2": "Akanksha Kulkarni",
        "head2_post": "Events Head"
    },
    "Development": {
        "head1": "Gouse Moideen S",
        "head1_post": "Technical Head",
        "head2": "Tarang Gupta",
        "head2_post": "Projects Head"
    },
    "Technical": {
        "head1": "Gouse Moideen S",
        "head1_post": "Technical Head",
        "head2": "Tarang Gupta",
        "head2_post": "Projects Head"
    },
    "AI/ML": {
        "head1": "Gouse Moideen S",
        "head1_post": "Technical Head",
        "head2": "Tarang Gupta",
        "head2_post": "Projects Head"
    },
    "Competitive Coding": {
        "head1": "Gouse Moideen S",
        "head1_post": "Technical Head",
        "head2": "Tarang Gupta",
        "head2_post": "Projects Head"
    },
    "Cyber Security": {
        "head1": "Gouse Moideen S",
        "head1_post": "Technical Head",
        "head2": "Tarang Gupta",
        "head2_post": "Projects Head"
    },
    "UI/UX": {
        "head1": "Gouse Moideen S",
        "head1_post": "Technical Head",
        "head2": "Tarang Gupta",
        "head2_post": "Projects Head"
    },
    "Design": {
        "head1": "Preeti B R",
        "head1_post": "Creatives Head",
        "head2": "Ahmed Sajjad Shihab",
        "head2_post": "Publicity Head"
    },
    "Creatives": {
        "head1": "Preeti B R",
        "head1_post": "Creatives Head",
        "head2": "Ahmed Sajjad Shihab",
        "head2_post": "Publicity Head"
    },
    "Social Media & Content": {
        "head1": "Preeti B R",
        "head1_post": "Creatives Head",
        "head2": "Ahmed Sajjad Shihab",
        "head2_post": "Publicity Head"
    }
}

position_mapping = {
    "Management": "Member of Management",
    "Entrepreneurship": "Member of Entrepreneurship",
    "Development": "Member of Development",
    "Technical": "Member of Technical",
    "Competitive Coding": "Member of Competitive Coding",
    "UI/UX": "Member of UI/UX",
    "Cyber Security": "Member of Cyber Security",
    "AI/ML": "Member of AI/ML",
    "Social Media & Content": "Member of Social Media & Content",
    "Creatives": "Member of Creatives",
    "Design": "Member of Design"
}

# =====================================================
# TEXT REPLACEMENT FUNCTION (RUN-AWARE & SPLIT-SAFE)
# =====================================================
def replace_text_in_paragraph(paragraph, key, value):
    """
    Safely replace placeholder text in a paragraph.
    Handles placeholders residing in single runs or split across multiple runs.
    """
    if key not in paragraph.text:
        return

    # First check if the full key is present inside any single run
    found_in_single_run = False
    for run in paragraph.runs:
        if key in run.text:
            run.text = run.text.replace(key, value)
            found_in_single_run = True

    # If key was split across multiple runs
    if not found_in_single_run and key in paragraph.text:
        # Combine text from all runs, replace, and place in the first run while clearing subsequent runs
        full_text = "".join(run.text for run in paragraph.runs)
        updated_text = full_text.replace(key, value)
        for i, run in enumerate(paragraph.runs):
            if i == 0:
                run.text = updated_text
            else:
                run.text = ""

def apply_replacements(doc, replacements):
    # Body paragraphs
    for p in doc.paragraphs:
        for k, v in replacements.items():
            replace_text_in_paragraph(p, k, v)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for k, v in replacements.items():
                        replace_text_in_paragraph(p, k, v)

    # Headers & Footers
    for section in doc.sections:
        for p in section.header.paragraphs:
            for k, v in replacements.items():
                replace_text_in_paragraph(p, k, v)
        for p in section.footer.paragraphs:
            for k, v in replacements.items():
                replace_text_in_paragraph(p, k, v)

# =====================================================
# MAIN GENERATION LOOP
# =====================================================
success_count = 0
error_count = 0

for index, row in df.iterrows():
    try:
        # Extract row details
        name = get_column_value(row, ["NAME", "Full Name", "Full_Name", "Member Name"])
        regno = get_column_value(row, ["REGISTRATION_NUMBER", "Registration Number", "Reg_No", "Registration_No", "RegNo"])
        dept = get_column_value(row, ["DEPARTMENT", "Department", "Dept", "Domain"])
        email = get_column_value(row, ["EMAIL_ID", "Email", "Email ID", "Email_ID"])

        if not name and not regno:
            continue

        print(f"\n[{index + 1}/{len(df)}] Processing: {name} | {regno} | {dept}")

        # Choose template
        if dept == "Entrepreneurship" and os.path.exists("MIC_Letter_Entrepreneurship.docx"):
            template_path = "MIC_Letter_Entrepreneurship.docx"
        elif os.path.exists(DEFAULT_TEMPLATE):
            template_path = DEFAULT_TEMPLATE
        else:
            raise FileNotFoundError(f"Template '{DEFAULT_TEMPLATE}' not found!")

        # Load fresh template
        doc = Document(template_path)

        # Determine position
        position = position_mapping.get(dept, f"Member of {dept}")

        # Get heads info
        heads = dept_heads.get(
            dept,
            {
                "head1": "",
                "head1_post": "",
                "head2": "",
                "head2_post": ""
            }
        )

        # Prepare replacement dictionary
        replacements = {
            "[Name]": name,
            "[NAME]": name,
            "{{NAME}}": name,
            "[Position]": position,
            "[POSITION]": position,
            "{{POSITION}}": position,
            "[Registration Number]": regno,
            "[REGISTRATION_NUMBER]": regno,
            "{{REGISTRATION_NUMBER}}": regno,
            "[Department]": dept,
            "[DEPARTMENT]": dept,
            "{{DEPARTMENT}}": dept,
            "[Email]": email,
            "[EMAIL_ID]": email,
            "{{EMAIL_ID}}": email,
            "[Head1]": heads["head1"],
            "[Head1_Post]": heads["head1_post"],
            "[Head2]": heads["head2"],
            "[Head2_Post]": heads["head2_post"]
        }

        # Apply replacements
        apply_replacements(doc, replacements)

        # Generate filenames
        dept_clean = (
            dept.lower()
            .replace("/", "")
            .replace("&", "and")
            .replace(" ", "_")
        ) if dept else "general"

        regno_clean = regno.replace("/", "_").replace(" ", "_") if regno else f"row_{index+1}"

        docx_filename = f"{dept_clean}_{regno_clean}.docx"
        pdf_filename = f"{dept_clean}_{regno_clean}.pdf"

        docx_path = os.path.join(DOCX_OUTPUT_DIR, docx_filename)
        pdf_path = os.path.join(PDF_OUTPUT_DIR, pdf_filename)

        # Save generated DOCX
        doc.save(docx_path)
        print(f" -> DOCX generated: {docx_filename}")

        # Convert to PDF
        try:
            convert(docx_path, pdf_path)

            # Keep only the first page if multiple pages were generated
            with open(pdf_path, "rb") as f:
                pdf_bytes = io.BytesIO(f.read())

            reader = PdfReader(pdf_bytes)
            if len(reader.pages) > 0:
                writer = PdfWriter()
                writer.add_page(reader.pages[0])

                with open(pdf_path, "wb") as output_pdf:
                    writer.write(output_pdf)

            print(f" -> PDF generated:  {pdf_filename}")
        except Exception as pdf_err:
            print(f" -> PDF Warning: {pdf_err}")

        success_count += 1

    except Exception as e:
        error_count += 1
        print(f"ERROR processing row {index + 1}: {e}")

print("\n" + "=" * 60)
print(f"COMPLETED! Successfully processed: {success_count}, Errors: {error_count}")
print(f"DOCX files saved to: ./{DOCX_OUTPUT_DIR}/")
print(f"PDF files saved to:  ./{PDF_OUTPUT_DIR}/")
print("=" * 60)
